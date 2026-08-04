import docx
import os

folder = r"E:\Nilal_thiruvila\SmartTriage_Dashboard\Niral_work_offline"
files = [
    "NMNT 2026-BILL SUMMARY FORMAT.docx",
    "NMNT 2026-PRE RECEIPT FORMAT.docx",
    "NMNT 2026-UTILISATION CERTIFICATE FORMAT.docx"
]

for filename in files:
    print("========================================")
    print(f"FILE: {filename}")
    print("========================================")
    try:
        doc = docx.Document(os.path.join(folder, filename))
        for para in doc.paragraphs:
            print(para.text)
        for table in doc.tables:
            print("\n--- TABLE ---")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.replace("\n", " ").strip())
                print(" | ".join(row_data))
            print("-------------\n")
    except Exception as e:
        print(f"Error reading {filename}: {e}")
    print("\n\n")
